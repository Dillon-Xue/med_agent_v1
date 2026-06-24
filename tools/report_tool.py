import os
import re
from datetime import datetime
from docx import Document
from utils.response import build_response

class ReportTool:
    def __init__(self):
        self.template_path = "templates/登记表模板.docx"
        self.output_dir = "reports/"
        os.makedirs(self.output_dir, exist_ok=True)

    def _log(self, msg):
        print(f"[ReportTool] {msg}")

    def run(self, query: str) -> dict:
        """统一入口：生成评估表"""
        clean_query = re.sub(r'^对话历史：.*?当前问题：', '', query)

        # ---- 优先从 query 中提取患者姓名 ----
        patient_name = None
        
        # 1. 提取“生成评估表 张三”或“张三 生成评估表”格式
        match = re.search(r'(?:生成评估表|生成病历|生成档案|生成记录)\s*([\u4e00-\u9fa5]{2,4})', clean_query)
        if not match:
            match = re.search(r'([\u4e00-\u9fa5]{2,4})\s*(?:生成评估表|生成病历|生成档案|生成记录)', clean_query)
        if not match:
            match = re.search(r'([\u4e00-\u9fa5]{2,4})\s*的(?:评估表|病历|档案)', clean_query)
        if match:
            patient_name = match.group(1)
            print(f"[ReportTool] 从 query 提取患者姓名: {patient_name}")

        # ---- 如果 query 中没有，再尝试从全局会话获取（兜底） ----
        if not patient_name:
            try:
                import chat
                if hasattr(chat, 'current_session_user') and chat.current_session_user:
                    patient_name = chat.current_session_user
                    print(f"[ReportTool] 从会话兜底获取患者姓名: {patient_name}")
            except Exception as e:
                print(f"[ReportTool] 会话兜底失败: {e}")

        # ---- 如果提取到的姓名是无效值，过滤掉 ----
        invalid_names = ["患者信息", "患者", "信息", "评估表", "生成评估表", "生成报告", "生成档案", "生成病历"]
        if patient_name in invalid_names:
            print(f"[ReportTool] 提取到无效姓名 '{patient_name}'，设为空")
            patient_name = None

        # ---- 如果还是没有，返回错误 ----
        if not patient_name:
            return build_response(
                answer="❌ 请指定患者姓名，例如：生成评估表 张三",
                source="report",
                success=False
            )

        # ---- 查询患者档案 ----
        from tools.tool_registry import get_tools
        tools = get_tools()
        patient_tool = tools.get("patient")
        candidates = patient_tool.search_patients(patient_name)

        if not candidates:
            return build_response(
                answer=f"❌ 未找到患者 {patient_name} 的档案，请先录入患者信息。",
                source="report",
                success=False
            )

        if len(candidates) == 1:
            return self._generate_from_candidate(candidates[0])

        return self._show_candidates(candidates, patient_name)

    def _show_candidates(self, candidates: list, name: str) -> dict:
        lines = [f"找到 {len(candidates)} 位名为 {name} 的患者，请选择："]
        for i, c in enumerate(candidates, 1):
            info_preview = c.get("info", "")[:40]
            lines.append(f"{i}. {c['name']}，{info_preview}...")
        lines.append("请回复数字（如 1）选择，或输入更详细的信息确认。")
        return build_response(
            answer="\n".join(lines),
            source="report",
            success=False,
            debug={"candidates": candidates}
        )

    def _generate_from_candidate(self, candidate: dict) -> dict:
        id_card = candidate.get("id_card", "")
        if not id_card:
            return build_response(
                answer=f"❌ 患者 {candidate['name']} 缺少身份证号，请先补充：\n记住患者 {candidate['name']}：身份证号 410123199001011234\n然后重新生成评估表。",
                source="report",
                success=False
            )

        # ===== 🆕 直接从结构化字段读取，不再用正则解析 =====
        info_dict = {
            "姓名": candidate.get("name", ""),
            "性别": candidate.get("gender", ""),
            "年龄": candidate.get("age", ""),
            "联系方式": candidate.get("phone", ""),
            "家庭住址": candidate.get("address", ""),
            "目前用药": candidate.get("medication", ""),
            "临床诊断": candidate.get("diagnosis", ""),
            "症状": candidate.get("symptoms", ""),
            "过敏史": candidate.get("allergy", ""),
            "主要问题": candidate.get("diagnosis", "")  # 默认用诊断作为主要问题
        }

        # 🆕 如果结构化字段都为空，降级用正则解析 info（兼容旧数据）
        if not info_dict["性别"] and not info_dict["年龄"] and not info_dict["联系方式"]:
            print(f"[ReportTool] 结构化字段为空，降级使用正则解析 info")
            info_dict = self._parse_patient_info(
                candidate.get("info", ""),
                candidate.get("diagnosis", "")
            )
        else:
            # 如果临床诊断为空，用症状填充
            if not info_dict["临床诊断"] and info_dict.get("症状"):
                info_dict["临床诊断"] = info_dict["症状"]
            # 如果主要问题为空，用临床诊断填充
            if not info_dict["主要问题"] and info_dict.get("临床诊断"):
                info_dict["主要问题"] = info_dict["临床诊断"]

        print(f"[ReportTool] 最终 info_dict: {info_dict}")

        if info_dict.get("家庭住址"):
            info_dict["家庭住址"] = self._complete_address(info_dict["家庭住址"])

        drug_suggestion = self._get_drug_suggestion(candidate["name"], info_dict)
        assessment = self._generate_assessment(candidate["name"], info_dict, drug_suggestion)

        output_path = self._fill_template(
            candidate["name"],
            info_dict,
            assessment,
            id_card
        )

        download_url = f"/reports/{os.path.basename(output_path)}"

        # ===== 🆕 自动创建审批项 =====
        try:
            from tools.approval_tool import ApprovalTool
            from chat import current_session_user

            approval_tool = ApprovalTool()
            requester = current_session_user if current_session_user else "system"
            reviewer = "doctor_张"

            # 构建审批内容
            content = f"""
    患者姓名：{candidate['name']}
    年龄：{info_dict.get('年龄', '未知')}
    性别：{info_dict.get('性别', '未知')}
    临床诊断：{info_dict.get('临床诊断', '无')}
    目前用药：{info_dict.get('目前用药', '无')}

    评估结果：{assessment.get('评估结果', '')}
    用药目标：{assessment.get('用药目标', '')}
    用药注意事项：{assessment.get('用药注意事项', '')}

    评估表下载：{download_url}
    """
            approval_tool.create(
                title=f"用药方案评估审批：{candidate['name']}",
                content=content,
                type="medication_evaluation",
                requester=requester,
                reviewer=reviewer
            )
            print(f"[ReportTool] 已自动创建审批项（审批人：{reviewer}）")
        except Exception as e:
            print(f"[ReportTool] 创建审批项失败: {e}")
            # 不影响主流程，继续返回下载链接

        return build_response(
            answer=f"✅ 评估表已生成：{candidate['name']}\n📎 [下载评估表]({download_url})",
            source="report",
            debug={
                "patient": candidate["name"],
                "output": output_path,
                "download_url": download_url
            }
        )

    def _complete_address(self, address: str) -> str:
        if not address:
            return address
        if re.search(r'省.*市', address):
            return address
        try:
            from openai import OpenAI
            client = OpenAI(
                api_key=os.getenv("DASHSCOPE_API_KEY"),
                base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
            )
            prompt = f"""
请将以下中国城市地址补全为完整的“省份+城市”格式，只输出结果，不要添加任何解释。
输入地址：{address}
输出格式：省份+城市
示例：长沙市 → 湖南省长沙市
"""
            resp = client.chat.completions.create(
                model="qwen-plus",
                messages=[{"role": "user", "content": prompt}],
                temperature=0
            )
            full_address = resp.choices[0].message.content.strip()
            if full_address:
                return full_address
        except Exception as e:
            print(f"[ReportTool] 地址补全失败: {e}")
        return address

    def _get_drug_suggestion(self, name: str, info: dict) -> str:
        try:
            from tools.tool_registry import get_tools
            tools = get_tools()
            drug_tool = tools.get("drug")
            if not drug_tool:
                return ""
            query = f"{name}，{info.get('临床诊断', '')}，{info.get('主要问题', '')}，请推荐适合的药品"
            result = drug_tool.run(query)
            if result and isinstance(result, dict):
                return result.get("answer", "")
        except Exception as e:
            print(f"[ReportTool] 调用 drug tool 失败: {e}")
        return ""

    def _parse_patient_info(self, text: str, diagnosis: str = "", 
                            structured_data: dict = None) -> dict:
        """
        解析患者信息
        优先使用 structured_data（结构化字段），否则从 text 中正则提取
        """
        info = {
            "姓名": "",
            "性别": "",
            "年龄": "",
            "联系方式": "",
            "家庭住址": "",
            "目前用药": "",
            "临床诊断": diagnosis or "",
            "主要问题": "",
        }

        # ===== 优先使用结构化数据 =====
        if structured_data:
            info["姓名"] = structured_data.get("name", "")
            info["性别"] = structured_data.get("gender", "")
            info["年龄"] = structured_data.get("age", "")
            info["联系方式"] = structured_data.get("phone", "")
            info["家庭住址"] = structured_data.get("address", "")
            info["目前用药"] = structured_data.get("medication", "")
            info["临床诊断"] = structured_data.get("diagnosis", "") or diagnosis
            
            # 如果临床诊断为空，用症状拼接
            if not info["临床诊断"] and structured_data.get("symptoms"):
                info["临床诊断"] = structured_data.get("symptoms", "")
            
            # 主要问题
            if info["临床诊断"] and not info["主要问题"]:
                info["主要问题"] = info["临床诊断"]
            
            return info

        # ===== 如果没有结构化数据，降级到正则提取 =====
        # ... 原有正则逻辑 ...
        return info

    def _generate_assessment(self, name: str, info: dict, drug_suggestion: str) -> dict:
        from openai import OpenAI
        client = OpenAI(
            api_key=os.getenv("DASHSCOPE_API_KEY"),
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )

        drug_reference = ""
        if drug_suggestion:
            drug_reference = f"【药物参考信息】\n{drug_suggestion}\n"

        prompt = f"""
请根据以下患者信息生成评估报告内容。

患者姓名：{name}
年龄：{info.get('年龄', '未知')}
性别：{info.get('性别', '未知')}
临床诊断：{info.get('临床诊断', '无')}
目前用药：{info.get('目前用药', '无')}
{drug_reference}
请生成以下四项内容，每项内容要简洁、专业，不要包含任何"助手"、"工具"、"来源"等标记。

1. 评估结果：简要概括患者当前的整体健康状况（40-60字）
2. 主要问题：从诊断中提取最核心的一个问题（20-30字）
3. 用药目标：根据患者诊断和目前用药，给出具体的用药目标和推荐方案。如果已有用药，写明继续或调整建议；如果无用药，结合药物参考信息给出具体药物推荐（包括药名和预期效果）。
4. 用药注意事项：根据患者情况，提醒用药过程中的关键注意事项（40-60字）

输出格式（每项占一行）：
评估结果：...
主要问题：...
用药目标：...
用药注意事项：...
"""
        resp = client.chat.completions.create(
            model="qwen-plus",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        content = resp.choices[0].message.content
        result = {"评估结果": "", "主要问题": "", "用药目标": "", "用药注意事项": ""}
        for line in content.split("\n"):
            if "评估结果：" in line:
                result["评估结果"] = line.replace("评估结果：", "").strip()
            elif "主要问题：" in line:
                result["主要问题"] = line.replace("主要问题：", "").strip()
            elif "用药目标：" in line:
                result["用药目标"] = line.replace("用药目标：", "").strip()
            elif "用药注意事项：" in line:
                result["用药注意事项"] = line.replace("用药注意事项：", "").strip()
        return result

    def _fill_template(self, name: str, info: dict, assessment: dict, id_card: str) -> str:
        doc = Document(self.template_path)
        now = datetime.now()
        档案号 = f"DA{now.strftime('%Y%m%d%H%M%S')}"

        # 如果目前用药为空，设为“无”
        if not info.get("目前用药"):
            info["目前用药"] = "无"

        raw_map = {
            "档案号": 档案号,
            "姓名": name,
            "身份证号": id_card,
            "家庭住址": info.get("家庭住址", ""),
            "评估时间": now.strftime("%Y年%m月%d日"),
            "性别": info.get("性别", ""),
            "年龄": info.get("年龄", ""),
            "联系方式": info.get("联系方式", ""),
            "临床诊断": info.get("临床诊断", ""),
            "主要问题": assessment.get("主要问题", ""),
            "目前用药": info.get("目前用药", ""),
            "评估结果": assessment.get("评估结果", ""),
            "用药目标": assessment.get("用药目标", ""),
            "用药注意事项": assessment.get("用药注意事项", ""),
        }

        replace_map = {}
        for key, value in raw_map.items():
            replace_map[f"{{{{{key}}}}}"] = value
            replace_map[f"《{key}》"] = value

        for para in doc.paragraphs:
            for placeholder, value in replace_map.items():
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in replace_map.items():
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, value)

        timestamp = now.strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(self.output_dir, f"{name}_评估表_{timestamp}.docx")
        doc.save(output_path)
        return output_path