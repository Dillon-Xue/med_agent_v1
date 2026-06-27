import os, base64, tempfile, logging
from typing import Optional
from openai import OpenAI
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from openpyxl import load_workbook
from utils.config import get_llm_client
logger = logging.getLogger(__name__)

class FileTool:
    """文件解析工具：支持图片（PNG/JPG）和 PDF 的内容提取"""
    
    def __init__(self, api_key: str):
        # 🆕 使用统一客户端工厂
        self.client, self.model = get_llm_client(api_key)
        logger.info(f"[FileTool] Using model: {self.model}")

    def parse_image(self, file_content: bytes, filename: str) -> str:
        """
        调用 Qwen-VL 识别图片内容
        """
        base64_image = base64.b64encode(file_content).decode('utf-8')
        
        ext = filename.split('.')[-1].lower()
        mime_type = "image/png" if ext == "png" else "image/jpeg"
        
        try:
            response = self.client.chat.completions.create(
                model="qwen-vl-plus",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{mime_type};base64,{base64_image}"}
                            },
                            {
                                "type": "text",
                                "text": (
                                    "请识别这张图片中的所有文字信息，并按照以下格式输出（如果某个字段不存在，请标注为'未提及'）：\n\n"
                                    "患者姓名：xxx\n"
                                    "年龄：xxx岁\n"
                                    "诊断/疾病名称：xxx\n"
                                    "药品名称和用法用量：xxx\n"
                                    "过敏史：xxx\n"
                                    "身份证号：xxx\n"
                                    "联系电话：xxx\n"
                                    "家庭住址：xxx\n\n"
                                    "然后请用自然语言总结这段医疗信息，重点关注与诊断、用药、过敏史相关的内容。"
                                )
                            }
                        ]
                    }
                ],
                temperature=0.1
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"图片识别失败: {str(e)}"

    def parse_pdf(self, file_content: bytes) -> str:
        """
        解析 PDF 文件，提取纯文本内容
        """
        tmp_path = None
        try:
            # 保存临时文件（PyPDFLoader 需要文件路径）
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            loader = PyPDFLoader(tmp_path)
            docs = loader.load()
            
            full_text = "\n".join([doc.page_content for doc in docs])
            # 如果提取到的文本为空，返回友好提示
            if not full_text or not full_text.strip():
                return "PDF 解析成功，但未提取到文本内容（可能是扫描件，建议上传图片格式）"
            
            # 截断防止上下文爆炸（保留前 15000 字符）
            if len(full_text) > 15000:
                full_text = full_text[:15000] + "\n...(PDF 内容过长，已截断)"
            return full_text
        except Exception as e:
            return f"PDF 解析失败: {str(e)}"
        finally:
            # 确保临时文件被删除（无论是否发生异常）
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass  # 忽略删除失败（权限问题等）

    def parse_docx(self, file_content: bytes, filename: str) -> str:
        """
        解析 DOCX 文件，提取段落和表格中的文本内容
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            os.unlink(tmp_path)
            
            text_parts = []
            
            # 1. 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            # 2. 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
            
            full_text = "\n".join(text_parts)
            
            # 如果仍然没有内容，可能是文档包含文本框或形状，尝试提取
            if not full_text:
                # 尝试使用 python-docx 的 element 方式提取所有文本（备用）
                try:
                    from docx.oxml import parse_xml
                    # 简单方法：遍历所有段落和运行
                    for para in doc.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                text_parts.append(run.text.strip())
                    full_text = "\n".join(text_parts)
                except:
                    pass
            
            if len(full_text) > 15000:
                full_text = full_text[:15000] + "\n...(文档内容过长，已截断)"
            
            return full_text if full_text else "（文档中未提取到文字内容，请确认文档包含可识别的文本）"
        except Exception as e:
            return f"DOCX 解析失败: {str(e)}"

    def parse_xlsx(self, file_content: bytes, filename: str) -> str:
        """
        解析 XLSX 文件，提取表格内容为文本
        """
        import tempfile
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            wb = load_workbook(tmp_path, data_only=True)
            os.unlink(tmp_path)
            
            text_parts = []
            for sheet in wb.worksheets:
                sheet_name = sheet.title
                text_parts.append(f"【工作表：{sheet_name}】")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            if len(full_text) > 15000:
                full_text = full_text[:15000] + "\n...(表格内容过长，已截断)"
            return full_text
        except Exception as e:
            return f"XLSX 解析失败: {str(e)}"

    def run(self, file_content: bytes, filename: str, file_type: str) -> dict:
        """
        统一入口：根据文件类型调用不同解析器
        """
        ext = filename.split('.')[-1].lower() if '.' in filename else ''

        if ext in ["png", "jpg", "jpeg"]:
            content = self.parse_image(file_content, filename)
        elif ext == "pdf":
            content = self.parse_pdf(file_content)
        elif ext == "docx":
            content = self.parse_docx(file_content, filename)
        elif ext in ["xlsx", "xls"]:
            content = self.parse_xlsx(file_content, filename)
        else:
            return {
                "success": False,
                "answer": f"不支持的文件格式: {ext}，目前支持 PNG、JPG、PDF、DOCX、XLSX",
                "source": "file"
            }

        return {
            "success": True,
            "answer": content,
            "source": "file",
            "filename": filename
        }