import os, re, json, hashlib, time, logging, pymysql, asyncio, httpx, requests, io
from pathlib import Path
from fastapi import FastAPI, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel
from logging.handlers import RotatingFileHandler
from contextvars import ContextVar
from typing import List, Optional
from agents.planner import Planner
from agents.executor import Executor
from agents.synthesizer import Synthesizer
from agents.consult_graph import ConsultGraph
from tools.tool_registry import get_tools
from docx import Document
from utils.audit import log_audit
from agents.supervisor import Supervisor  
from agents.agent_factory import get_agent
from agents.aggregator import Aggregator
from utils.config import setup_logging
from utils.database import get_connection
from tools.memory_tool import MemoryTool
logger = logging.getLogger(__name__)


# =========================
# Trace 存储（内存）
# =========================
trace_store = {}
trace_lock = asyncio.Lock()

# =========================
# 日志配置
# =========================
LOG_FILE = "logs/app.log"
os.makedirs("logs", exist_ok=True)

# =========================
# 路径配置
# =========================
BASE_DIR = Path(os.path.dirname(os.path.abspath(__file__)))
REPORTS_DIR = (BASE_DIR / "reports").resolve()

setup_logging()
logger = logging.getLogger("med_agent")
logger.setLevel(logging.INFO)

class TenantFilter(logging.Filter):
    def filter(self, record):
        try:
            from chat import get_current_tenant
            record.tenant_id = get_current_tenant()
        except:
            record.tenant_id = "default"
        return True

console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
file_handler = RotatingFileHandler(
    LOG_FILE,
    maxBytes=10*1024*1024,
    backupCount=5
)
file_handler.setLevel(logging.INFO)

formatter = logging.Formatter(
    '%(asctime)s - %(tenant_id)s - %(name)s - %(levelname)s - %(message)s'
)
console_handler.setFormatter(formatter)
file_handler.setFormatter(formatter)

logger.addFilter(TenantFilter())
logger.addHandler(console_handler)
logger.addHandler(file_handler)

# =========================
# 响应模型（用于 Swagger 文档）
# =========================
class ToolResult(BaseModel):
    success: bool
    id: str
    source: str
    answer: str
    debug: dict
    trace: list
    timestamp: float

class ResultData(BaseModel):
    answer: str
    tools_used: List[str]
    plan: dict
    tool_results: List[ToolResult]

class AskResponse(BaseModel):
    success: bool
    result: ResultData
    trace: dict

class ApprovalItem(BaseModel):
    id: str
    title: str
    requester: str
    created_at: str

class ApprovalsResponse(BaseModel):
    count: int
    items: List[ApprovalItem]

class UploadResponse(BaseModel):
    success: bool
    filename: str
    content: str
    preview: str          # 前 200 字符预览
    source: str
    error: Optional[str] = None

class HistoryItem(BaseModel):
    id: int
    role: str
    content: str
    tools_used: Optional[List[str]] = None
    file_name: Optional[str] = None
    conversation_type: str
    created_at: str

class HistoryResponse(BaseModel):
    success: bool
    items: List[HistoryItem]
    count: int

# =========================
# FastAPI App
# =========================
tags_metadata = [
    {"name": "问答", "description": "医学问答相关接口（快速问答、智能问诊）"},
    {"name": "审批", "description": "审批管理相关接口（列表、通过、驳回）"},
    {"name": "报告", "description": "报告生成与下载相关接口"},
    {"name": "运维", "description": "系统运维相关接口（健康检查、监控指标）"}
]

app = FastAPI(
    title="医疗Agent API",
    description="基于RAG的多工具医学问答助手，支持快速问答、智能问诊、评估表生成、多租户审批",
    version="3.0.0",
    openapi_tags=tags_metadata
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "http://localhost:3000,http://localhost:5173").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

# =========================
# 租户上下文
# =========================
tenant_id_var: ContextVar[str] = ContextVar("tenant_id", default="default")

def get_current_tenant() -> str:
    return tenant_id_var.get()

@app.middleware("http")
async def tenant_middleware(request: Request, call_next):
    tenant_id = request.headers.get("X-Tenant-ID", "default")
    token = tenant_id_var.set(tenant_id)
    try:
        response = await call_next(request)
        return response
    finally:
        tenant_id_var.reset(token)

# =========================
# 全局用户状态（用于审批/患者操作）
# =========================
current_session_user = None

# =========================
# 组件初始化
# =========================
TOOLS = get_tools()
supervisor = Supervisor()
aggregator = Aggregator()
planner = Planner()
executor = Executor(TOOLS)
synthesizer = Synthesizer(api_key=os.getenv("DASHSCOPE_API_KEY"))
consult_graph = ConsultGraph()

class ChatRequest(BaseModel):
    question: str
    history: list = []

# =========================
# 缓存
# =========================
class SimpleCache:
    def __init__(self):
        self._cache = {}
    def get(self, key):
        if key in self._cache:
            entry = self._cache[key]
            if time.time() - entry['time'] < entry['ttl']:
                return entry['value']
            else:
                del self._cache[key]
        return None
    def set(self, key, value, ttl=3600):
        self._cache[key] = {'value': value, 'time': time.time(), 'ttl': ttl}
    def clear(self):
        self._cache.clear()

cache = SimpleCache()

def should_cache(question: str, history: list) -> bool:
    pronouns = ["它", "他", "她", "这个", "那个", "刚才", "上述", "以上", "以下"]
    if any(word in question for word in pronouns):
        logger.info(f"[Cache] 包含指代词，不缓存: {question}")
        return False
    if not history or len(history) <= 1:
        return True
    return False

def get_cache_key(question: str) -> str:
    return hashlib.md5(question.encode()).hexdigest()

# =========================
# 核心问答逻辑
# =========================
async def process_question(question: str, history: list, trace_callback=None) -> dict:
    global current_session_user
    logger.info(f"[process_question] 收到的 trace_callback 是否为 None: {trace_callback is None}")
    
    def _is_topic_shift(current_q: str, history: list) -> bool:
        if not history:
            return False
        import re
        current_keywords = set(re.findall(r'[\u4e00-\u9fa5]{2,6}', current_q))
        user_messages = [msg["content"] for msg in history[-6:] if msg.get("role") == "user"]
        if not user_messages:
            return False
        history_keywords = set()
        for msg in user_messages:
            history_keywords.update(re.findall(r'[\u4e00-\u9fa5]{2,6}', msg))
        if not history_keywords or not current_keywords:
            return False
        overlap = len(current_keywords & history_keywords)
        overlap_rate = overlap / len(current_keywords)
        is_shift = overlap_rate < 0.3
        if is_shift:
            logger.info(f"[话题切换] 重叠度 {overlap_rate:.2%}，截断历史")
        return is_shift

    stripped = question.strip()
    greetings = ["你好", "您好", "hi", "hello", "在吗", "在不在", "你好呀"]
    if stripped in greetings or stripped.lower() in greetings:
        return {
            "success": True,
            "result": {
                "answer": "您好！我是医学问答助手，您可以向我咨询药物、指南、文献或风险相关问题。",
                "tools_used": [],
                "plan": {"question": question, "tools": []},
                "tool_results": []
            },
            "trace": {"executor": []}
        }

    # ==================== 1. 构建对话上下文 ====================
    if history:
        if _is_topic_shift(question, history):
            history = []
            logger.info("[话题切换] 检测到切换，清空历史")
        recent = history[-5:]
        filtered = []
        for msg in recent:
            if msg["role"] == "user" and msg["content"] in ["你好", "您好", "hi", "hello"]:
                continue
            filtered.append(msg)
        if filtered:
            history_str = "\n".join([f"{msg['role']}: {msg['content']}" for msg in filtered])
            augmented_question = f"对话历史：\n{history_str}\n当前问题：{question}"
        else:
            augmented_question = question
    else:
        augmented_question = question

    # ==================== 2. 加载患者档案 ====================
    patient_context = ""
    if "患者" in question and "记住" not in question and "记录" not in question and "追加" not in question and "补充" not in question:
        name_match = re.search(r'患者\s*([\u4e00-\u9fa5]{2,4})', question)
        if name_match:
            name = name_match.group(1)
            from tools.tool_registry import get_tools
            tools = get_tools()
            patient_tool = tools.get("patient")
            if patient_tool:
                info = patient_tool.recall(name)
                if info:
                    patient_context = f"【患者档案】姓名：{name}，{info}\n"
                    logger.info(f"[Patient] 已加载患者 {name} 档案")
        else:
            name_match = re.search(r'([\u4e00-\u9fa5]{2,4})\s*的?信息', question)
            if name_match and "记住" not in question and "追加" not in question:
                name = name_match.group(1)
                from tools.tool_registry import get_tools
                tools = get_tools()
                patient_tool = tools.get("patient")
                if patient_tool:
                    info = patient_tool.recall(name)
                    if info:
                        patient_context = f"【患者档案】姓名：{name}，{info}\n"
                        logger.info(f"[Patient] 已加载患者 {name} 档案")

    # ==================== 3. L4 语义记忆检索 ====================
    memory_context = ""
    try:
        from tools.memory_tool import MemoryTool
        doctor_id = current_session_user if current_session_user else None
        memory_tool = MemoryTool()
        results = memory_tool.recall(question, k=2, doctor_id=doctor_id, min_similarity=0.3)
        if results:
            memory_lines = ["【历史参考病例】"]
            for r in results:
                meta = r["metadata"]
                memory_lines.append(
                    f"- {meta.get('patient_name', '未知')}：{meta.get('diagnosis', '')}，"
                    f"用药：{meta.get('medications', '无')}，"
                    f"相似度：{(1 - r.get('score', 0)):.0%}"
                )
            memory_context = "\n".join(memory_lines) + "\n\n"
            logger.info(f"[Memory] 检索到 {len(results)} 条历史病例")
    except Exception as e:
        logger.error(f"[Memory] 检索失败: {e}")

    # ==================== 4. 合并最终上下文 ====================
    final_augmented = ""
    if patient_context:
        final_augmented += patient_context
    if memory_context:
        final_augmented += memory_context
    final_augmented += augmented_question

    # ==================== 5. 多 Agent 执行 ====================
    route = supervisor.route(question)
    primary = route.get("primary", "general")
    secondary = route.get("secondary", [])
    all_agents = [primary] + secondary

    if trace_callback:
        trace_callback("supervisor", {"primary": primary, "secondary": secondary})

    if len(all_agents) == 1:
        agent = get_agent(primary)
        result = await agent.run(final_augmented, history, trace_callback=trace_callback)
        final_answer = result["answer"]
        tool_list = result.get("tools_used", [])
        tool_results = result.get("tool_results", [])
    else:
        import asyncio
        agents = [get_agent(s) for s in all_agents]
        tasks = [a.run(final_augmented, history, trace_callback=trace_callback) for a in agents]
        results = await asyncio.gather(*tasks)
        agent_answers = {r["specialty"]: r["answer"] for r in results}
        tool_list = []
        tool_results = []
        for r in results:
            tool_list.extend(r.get("tools_used", []))
            tool_results.extend(r.get("tool_results", []))
        tool_list = list(dict.fromkeys(tool_list))
        final_answer = aggregator.run(question, agent_answers)

    plan = {"question": question, "tools": tool_list, "mode": "multi_agent"}
    return {
        "success": True,
        "result": {
            "answer": final_answer,
            "tools_used": tool_list,
            "plan": plan,
            "tool_results": tool_results
        },
        "trace": {"executor": tool_results}
    }

def save_conversation(session_id: str, role: str, content: str, 
                      tools_used: list = None, file_name: str = None,
                      conversation_type: str = "quick", tenant_id: str = None):
    """保存单条对话记录到 conversations 表"""
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        tools_json = json.dumps(tools_used) if tools_used else None
        if tenant_id is None:
            tenant_id = get_current_tenant()
        
        cursor.execute(
            """INSERT INTO conversations 
               (session_id, tenant_id, role, content, tools_used, file_name, conversation_type) 
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (session_id, tenant_id, role, content, tools_json, file_name, conversation_type)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"保存对话历史失败: {e}")


# =========================
# 快速问答端点 (/ask)
# =========================
@app.post(
    "/v1/ask",
    tags=["问答"],
    summary="快速问答（V1）",
    description="单轮快速问答，适用于明确、完整的医学问题。支持多工具协同（drug/guideline/literature/risk/patient）。",
    response_model=AskResponse
)
@app.post(
    "/ask",
    tags=["问答"],
    summary="快速问答（V1）",
    description="单轮快速问答，适用于明确、完整的医学问题。支持多工具协同（drug/guideline/literature/risk/patient）。",
    response_model=AskResponse
)
async def ask(req: ChatRequest, request: Request):
    import re 
    import asyncio
    import uuid
    logger.info(f"[ASK] ==== 请求开始 ====")
    logger.debug(f"[ASK] 所有 headers: {request.headers}")
    logger.info(f"[ASK] X-Trace-ID: {request.headers.get('X-Trace-ID')}")
    global current_session_user
    question = req.question.strip()
    history = req.history
    session_id = request.headers.get("X-Session-ID")
    
    # 1. 身份声明
    if re.match(r'^用户[：:]', question):
        user_match = re.search(r'用户[：:]\s*(\S+)', question)
        if user_match:
            current_session_user = user_match.group(1)
            logger.info(f"[身份声明] 当前用户设置为: {current_session_user}")
            return {
                "success": True,
                "result": {
                    "answer": f"✅ 已识别当前用户：{current_session_user}",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }
    
    # 2. 身份查询
    if question in ["用户是谁", "我是谁", "当前用户", "whoami"]:
        if current_session_user:
            return {
                "success": True,
                "result": {
                    "answer": f"👤 当前用户：{current_session_user}",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }
        else:
            return {
                "success": True,
                "result": {
                    "answer": "⚠️ 当前未识别用户，请先声明身份：用户：doctor_张",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }

    # 3. 审批指令拦截
    approval_keywords = ["待审批", "审批通过", "驳回", "已通过", "已驳回", "全部列表", "审批列表"]
    if any(kw in question for kw in approval_keywords):
        logger.info(f"[Chat] 拦截到审批指令，直接处理: {question}")
        session_id = request.headers.get("X-Session-ID")
        conv_type = request.headers.get("X-Conversation-Type", "unknown")
        logger.info(f"[Chat] 审批拦截 - session_id: {session_id}, conv_type: {conv_type}")

        from tools.approval_tool import ApprovalTool
        approval_tool = ApprovalTool()
        result = approval_tool.run(question)
        """
        try:
            from utils.audit import log_audit
            answer = result.get("answer", "")
            if "审批通过" in answer:
                import re
                match = re.search(r'审批\s*([A-Z0-9\-]+)\s*已通过', answer)
                if match:
                    approval_id = match.group(1)
                    log_audit(
                        action="APPROVE",
                        resource_type="approval",
                        resource_id=approval_id,
                        detail={"status": "approved"},
                        ip=request.client.host
                    )
            elif "已驳回" in answer:
                match = re.search(r'审批\s*([A-Z0-9\-]+)\s*已驳回', answer)
                if match:
                    approval_id = match.group(1)
                    log_audit(
                        action="REJECT",
                        resource_type="approval",
                        resource_id=approval_id,
                        detail={"status": "rejected"},
                        ip=request.client.host
                    )
        except Exception as e:
            logger.error(f"[Audit] 审批日志记录失败: {e}")
        """

        if session_id:
            logger.debug("[Chat] 审批拦截 - 开始保存对话")
            try:
                save_conversation(session_id, "user", question, conversation_type="approval")
                answer = result.get("answer", "")
                save_conversation(session_id, "assistant", answer, conversation_type="approval")
                logger.debug("[Chat] 审批拦截 - 保存完成")
            except Exception as e:
                logger.error(f"[Chat] 审批拦截 - 保存失败: {e}")
        else:
            logger.warning("[Chat] 审批拦截 - session_id 为空，跳过保存")
        return {
            "success": True,
            "result": {
                "answer": result.get("answer", ""),
                "tools_used": ["approval"],
                "plan": {"question": question, "tools": ["approval"]},
                "tool_results": [result]
            },
            "trace": {"executor": [result]}
        }

    # 4. 患者操作拦截
    if re.search(r'记住患者|记录患者|追加患者|补充患者', question, re.IGNORECASE):
        logger.info(f"[Chat] 拦截到患者操作: {question}")
        from tools.patient_tool import PatientTool
        patient_tool = PatientTool()
        name_match = re.search(r'(?:记住患者|记录患者|追加患者|补充患者)\s*([\u4e00-\u9fa5]{2,4})\s*[:：]?\s*(.+)', question)
        if name_match:
            name = name_match.group(1)
            info = name_match.group(2).strip()
            result = patient_tool.remember(name, info, append=False)
            return {
                "success": True,
                "result": {
                    "answer": result.get("answer", ""),
                    "tools_used": ["patient"],
                    "plan": {"question": question, "tools": ["patient"]},
                    "tool_results": [result]
                },
                "trace": {"executor": [result]}
            }
        else:
            return {
                "success": False,
                "result": {
                    "answer": "❌ 无法解析患者信息，请使用格式：记住患者 张三：60岁，男，肚子痛",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }

    # ============================================================
    # 5. 正常问答（错误处理 + 超时控制）
    # ============================================================
    cache_key = get_cache_key(question)
    trace_id = request.headers.get("X-Trace-ID") or request.headers.get("x-trace-id")
    logger.info(f"[Trace] 收到 trace_id: {trace_id}")

    def trace_callback(step_type: str, data: dict):
        logger.info(f"[Trace] 收到 step: {step_type}, data: {data}")
        if not trace_id:
            return
        try:
            logger.info(f"[Trace] 正在记录 step: {step_type}")
            if trace_id in trace_store:
                trace_store[trace_id]["steps"].append({
                    "step_type": step_type,
                    "timestamp": time.time(),
                    "data": data
                })
            logger.info(f"[Trace] step 记录完成: {step_type}")
        except Exception as e:
            logger.error(f"[Trace] 回调失败: {e}")

    # 🆕 启动 trace
    if trace_id:
        try:
            async with httpx.AsyncClient() as client:
                await client.post(
                    "http://localhost:8000/trace/start",
                    json={"session_id": trace_id, "question": question},
                    timeout=2.0
                )
            logger.debug(f"[Trace] 启动成功: {trace_id}")
        except Exception as e:
            logger.error(f"[Trace] 启动失败: {e}")

    # ===== 执行问答 =====
    try:
        # 超时控制：整个问答流程最多 60 秒
        result = await asyncio.wait_for(
            _execute_ask(question, history, trace_id, cache_key, trace_callback),
            timeout=60.0
        )
    except asyncio.TimeoutError:
        error_id = str(uuid.uuid4())[:8]
        logger.error(f"[Ask] 请求超时 (error_id={error_id})")
        return {
            "success": False,
            "result": {
                "answer": f"⏱️ 请求超时（60秒），请稍后重试。（错误ID：{error_id}）",
                "tools_used": [],
                "plan": {"question": question, "tools": []},
                "tool_results": []
            },
            "trace": {"executor": []}
        }
    except Exception as e:
        error_id = str(uuid.uuid4())[:8]
        logger.error(f"[Ask] 异常: {e} (error_id={error_id})", exc_info=True)
        return {
            "success": False,
            "result": {
                "answer": f"❌ 服务暂时不可用，请稍后重试。（错误ID：{error_id}）",
                "tools_used": [],
                "plan": {"question": question, "tools": []},
                "tool_results": []
            },
            "trace": {"executor": []}
        }

    # ===== 保存历史记录 =====
    session_id = request.headers.get("X-Session-ID")
    conv_type = request.headers.get("X-Conversation-Type", "quick")
    if session_id and result and result.get("success"):
        try:
            save_conversation(session_id, "user", question, conversation_type=conv_type)
            answer = result.get("result", {}).get("answer", "")
            tools_used = result.get("result", {}).get("tools_used", [])
            save_conversation(session_id, "assistant", answer, tools_used=tools_used, conversation_type=conv_type)
            logger.info(f"[历史记录] 已保存会话 {session_id}，类型：{conv_type}")
        except Exception as e:
            logger.error(f"[历史记录] 保存失败: {e}")

    # 更新 trace 结束时间
    if trace_id and trace_id in trace_store:
        trace_store[trace_id]["end_time"] = time.time()

    return result


# ============================================================
# 🆕 辅助函数：提取正常的问答执行逻辑
# ============================================================
async def _execute_ask(question: str, history: list, trace_id: str, cache_key: str, trace_callback):
    """将缓存判断和 process_question 调用独立出来，便于超时控制"""
    result = None

    if should_cache(question, history):
        cached_result = cache.get(cache_key)
        if cached_result:
            logger.info(f"[Cache] 命中缓存，问题：{question}")
            result = cached_result
            if trace_id:
                try:
                    if trace_id in trace_store:
                        trace_store[trace_id]["steps"].append({
                            "step_type": "cache",
                            "timestamp": time.time(),
                            "data": {"status": "hit", "question": question}
                        })
                except Exception as e:
                    logger.error(f"[Trace] 缓存命中记录失败: {e}")
        else:
            logger.info(f"[Cache] 未命中缓存，执行完整流程，问题：{question[:50]}")
            result = await process_question(question, history, trace_callback=trace_callback if trace_id else None)
            cache.set(cache_key, result, ttl=3600)
    else:
        logger.info(f"[Cache] 不适合缓存，执行完整流程，问题：{question[:50]}")
        result = await process_question(question, history, trace_callback=trace_callback if trace_id else None)

    return result

# =========================
# 智能问诊端点 (/consult)
# =========================
@app.post(
    "/consult",
    tags=["问答"],
    summary="智能问诊（V2）",
    description="基于 LangGraph 的多轮交互式问诊。Agent 会主动追问缺失信息（年龄、过敏史、用药史），收集完整后给出个性化用药建议。",
    response_model=AskResponse
)
async def consult(req: ChatRequest, request: Request):
    import asyncio
    import uuid
    global current_session_user
    question = req.question.strip()
    history = req.history

    # 身份声明
    if re.match(r'^用户[：:]', question):
        user_match = re.search(r'用户[：:]\s*(\S+)', question)
        if user_match:
            current_session_user = user_match.group(1)
            logger.info(f"[Consult] 身份声明，当前用户设置为: {current_session_user}")
            return {
                "success": True,
                "result": {
                    "answer": f"✅ 已识别当前用户：{current_session_user}",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }

    # 身份查询
    if question in ["用户是谁", "我是谁", "当前用户", "whoami"]:
        if current_session_user:
            return {
                "success": True,
                "result": {
                    "answer": f"👤 当前用户：{current_session_user}",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }
        else:
            return {
                "success": True,
                "result": {
                    "answer": "⚠️ 当前未识别用户，请先声明身份：用户：doctor_张",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }

    # 患者操作拦截
    if re.search(r'记住患者|记录患者|追加患者|补充患者', question, re.IGNORECASE):
        logger.info(f"[Consult] 拦截到患者操作: {question}")
        from tools.patient_tool import PatientTool
        patient_tool = PatientTool()
        name_match = re.search(r'(?:记住患者|记录患者|追加患者|补充患者)\s*([\u4e00-\u9fa5]{2,4})\s*[:：]?\s*(.+)', question)
        if name_match:
            name = name_match.group(1)
            info = name_match.group(2).strip()
            result = patient_tool.remember(name, info, append=False)
            return {
                "success": True,
                "result": {
                    "answer": result.get("answer", ""),
                    "tools_used": ["patient"],
                    "plan": {"question": question, "tools": ["patient"]},
                    "tool_results": [result]
                },
                "trace": {"executor": [result]}
            }
        else:
            return {
                "success": False,
                "result": {
                    "answer": "❌ 无法解析患者信息，请使用格式：记住患者 张三：60岁，男，肚子痛",
                    "tools_used": [],
                    "plan": {"question": question, "tools": []},
                    "tool_results": []
                },
                "trace": {"executor": []}
            }

    # 审批指令拦截
    approval_keywords = ["待审批", "审批通过", "驳回", "已通过", "已驳回", "全部列表", "审批列表"]
    if any(kw in question for kw in approval_keywords):
        logger.info(f"[Consult] 拦截到审批指令: {question}")
        from tools.approval_tool import ApprovalTool
        approval_tool = ApprovalTool()
        result = approval_tool.run(question)

        session_id = request.headers.get("X-Session-ID")
        if session_id:
            save_conversation(session_id, "user", question, conversation_type="approval")
            answer = result.get("answer", "")
            save_conversation(session_id, "assistant", answer, conversation_type="approval")

        return {
            "success": True,
            "result": {
                "answer": result.get("answer", ""),
                "tools_used": ["approval"],
                "plan": {"question": question, "tools": ["approval"]},
                "tool_results": [result]
            },
            "trace": {"executor": [result]}
        }

    # ============================================================
    # 正常进入 LangGraph 问诊（超时控制 + 错误处理）
    # ============================================================
    logger.info(f"[Consult] 正常进入 LangGraph 问诊流程")
    
    try:
        # 60秒超时，用 to_thread 避免阻塞事件循环
        answer = await asyncio.wait_for(
            asyncio.to_thread(consult_graph.run, question, history),
            timeout=60.0
        )
    except asyncio.TimeoutError:
        error_id = str(uuid.uuid4())[:8]
        logger.error(f"[Consult] 问诊超时 (error_id={error_id})")
        return {
            "success": False,
            "result": {
                "answer": f"⏱️ 智能问诊超时（60秒），请稍后重试。（错误ID：{error_id}）",
                "tools_used": [],
                "plan": {"question": question, "mode": "consult"},
                "tool_results": []
            },
            "trace": {"executor": []}
        }
    except Exception as e:
        error_id = str(uuid.uuid4())[:8]
        logger.error(f"[Consult] 问诊异常: {e} (error_id={error_id})", exc_info=True)
        return {
            "success": False,
            "result": {
                "answer": f"❌ 智能问诊服务暂时不可用，请稍后重试。（错误ID：{error_id}）",
                "tools_used": [],
                "plan": {"question": question, "mode": "consult"},
                "tool_results": []
            },
            "trace": {"executor": []}
        }

    # 保存会话历史
    session_id = request.headers.get("X-Session-ID")
    if session_id:
        try:
            save_conversation(session_id, "user", question, conversation_type="consult")
            save_conversation(session_id, "assistant", answer, conversation_type="consult")
            logger.info(f"[历史记录] 已保存智能问诊会话 {session_id}")
        except Exception as e:
            logger.error(f"[历史记录] 保存失败: {e}")

    return {
        "success": True,
        "result": {
            "answer": answer,
            "tools_used": [],
            "plan": {"question": question, "mode": "consult"},
            "tool_results": []
        },
        "trace": {"executor": []}
    }

# =========================
# 审批列表接口（用于前端侧边栏）
# =========================
@app.get(
    "/approvals",
    tags=["审批"],
    summary="获取待审批列表",
    description="查询当前用户的待审批项列表，供前端侧边栏轮询使用。支持多租户隔离，只返回当前租户的数据。",
    response_model=ApprovalsResponse
)
async def get_approvals():
    from tools.tool_registry import get_tools
    from chat import current_session_user

    logger.debug(f"[DEBUG] /approvals - current_session_user: {current_session_user}")
    user = current_session_user if current_session_user else "current_user"
    logger.debug(f"[DEBUG] /approvals - 查询用户: {user}")

    tools = get_tools()
    approval_tool = tools.get("approval")
    if approval_tool:
        items = approval_tool.list_pending_by_user(user)
        logger.debug(f"[DEBUG] /approvals - 查询到 {len(items)} 条记录")
        return {"count": len(items), "items": items}
    return {"count": 0, "items": []}

# =========================
# 健康检查
# =========================
@app.get(
    "/health",
    tags=["运维"],
    summary="健康检查",
    description="检查服务是否正常运行，返回固定状态 'ok'。"
)
async def health_check():
    return {"status": "ok"}

# =========================
# Prometheus 监控
# =========================
from prometheus_client import Counter, Histogram, generate_latest, REGISTRY

REQUEST_COUNT = Counter('http_requests_total', 'Total HTTP requests', ['method', 'endpoint', 'tenant_id'])
REQUEST_LATENCY = Histogram('http_request_duration_seconds', 'HTTP request latency', ['method', 'endpoint', 'tenant_id'])
ERROR_COUNT = Counter('http_errors_total', 'Total HTTP errors', ['method', 'endpoint', 'status', 'tenant_id'])

@app.middleware("http")
async def metrics_middleware(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    latency = time.time() - start_time
    tenant_id = get_current_tenant()
    REQUEST_COUNT.labels(
        method=request.method,
        endpoint=request.url.path,
        tenant_id=tenant_id
    ).inc()
    REQUEST_LATENCY.labels(
        method=request.method,
        endpoint=request.url.path,
        tenant_id=tenant_id
    ).observe(latency)
    if response.status_code >= 400:
        ERROR_COUNT.labels(
            method=request.method,
            endpoint=request.url.path,
            status=response.status_code,
            tenant_id=tenant_id
        ).inc()
    return response

@app.get(
    "/metrics",
    tags=["运维"],
    summary="Prometheus 监控指标",
    description="暴露 Prometheus 格式的监控指标（请求数、延迟、错误数等），支持租户维度拆分。"
)
async def metrics():
    return Response(generate_latest(REGISTRY), media_type="text/plain")

# =========================
# 报告下载
# =========================
@app.get(
    "/reports/{filename}",
    tags=["报告"],
    summary="下载评估表",
    description="根据文件名下载生成的评估表 Word 文档（.docx 格式）。",
    responses={
        200: {"description": "文件下载成功", "content": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document": {}}},
        404: {"description": "文件不存在"}
    }
)
async def download_report(filename: str):
    # 安全校验：只取文件名部分，防止路径遍历
    safe_filename = Path(filename).name
    file_path = (REPORTS_DIR / safe_filename).resolve()

    # 确保最终路径在 reports 目录内
    if not str(file_path).startswith(str(REPORTS_DIR)):
        return {"error": "非法路径"}

    if not file_path.exists():
        return {"error": "文件不存在"}

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_filename
    )


from fastapi import File, UploadFile, Form
from typing import Optional
from tools.file_tool import FileTool

# 在 app 初始化后，初始化 FileTool
file_tool = FileTool(api_key=os.getenv("DASHSCOPE_API_KEY"))

@app.post(
    "/upload",
    tags=["文件"],
    summary="上传并解析文件（图片/PDF）",
    description="上传图片或 PDF 文件，系统自动识别并提取其中的医疗信息（患者姓名、年龄、诊断、药品等），返回纯文本内容供后续对话使用。",
    response_model=UploadResponse
)
async def upload_file(
    file: UploadFile = File(..., description="支持 PNG、JPG、PDF 格式，最大 10MB"),
    module: str = Form("consult", description="调用模块：quick | consult | approval")
):
    """
    上传文件并解析内容
    
    - **图片 (PNG/JPG)**：调用 Qwen-VL 多模态模型识别文字和医疗信息
    - **PDF**：提取全文文本，自动截断超长内容
    """
    # 1. 文件大小校验（10MB）
    content_bytes = await file.read()
    if len(content_bytes) > 10 * 1024 * 1024:
        return UploadResponse(
            success=False,
            filename=file.filename,
            content="",
            preview="",
            source="file",
            error="文件大小超过 10MB 限制"
        )
    
    # 2. 调用 FileTool 解析
    result = file_tool.run(content_bytes, file.filename, module)
    
    if not result.get("success"):
        return UploadResponse(
            success=False,
            filename=file.filename,
            content="",
            preview="",
            source="file",
            error=result.get("answer", "解析失败")
        )
    
    content = result.get("answer", "")
    preview = content[:200] + ("..." if len(content) > 200 else "")
    
    return UploadResponse(
        success=True,
        filename=file.filename,
        content=content,
        preview=preview,
        source="file"
    )

@app.get(
    "/history",
    tags=["会话"],
    summary="获取会话历史",
    description="根据 session_id 拉取最近 50 条对话记录，用于刷新页面后恢复聊天上下文",
    response_model=HistoryResponse
)
async def get_history(session_id: str, limit: int = 50, conversation_type: str = None):
    """
    获取指定会话的历史记录

    - **session_id**: 会话标识（前端生成，存储在 localStorage）
    - **limit**: 返回条数，默认 50
    - **conversation_type**: 可选过滤（quick / consult / approval）
    """
    try:
        conn = get_connection()
        cursor = conn.cursor()

        sql = """SELECT id, role, content, tools_used, file_name, conversation_type, created_at
                 FROM conversations
                 WHERE session_id = %s AND tenant_id = %s"""
        params = [session_id, get_current_tenant()]

        if conversation_type:
            sql += " AND conversation_type = %s"
            params.append(conversation_type)

        sql += " ORDER BY created_at ASC LIMIT %s"
        params.append(limit)

        cursor.execute(sql, params)
        rows = cursor.fetchall()
        conn.close()

        items = []
        for row in rows:
            tools = json.loads(row.get('tools_used')) if row.get('tools_used') else []
            items.append({
                "id": row.get('id'),
                "role": row.get('role'),
                "content": row.get('content'),
                "tools_used": tools,
                "file_name": row.get('file_name'),
                "conversation_type": row.get('conversation_type'),
                "created_at": row.get('created_at').strftime("%Y-%m-%d %H:%M:%S") if row.get('created_at') else ""
            })

        return {"success": True, "items": items, "count": len(items)}
    except Exception as e:
        logger.error(f"获取会话历史失败: {e}")
        return {"success": False, "items": [], "count": 0, "error": str(e)}

# =========================
# Trace 存储（内存，生产环境可改为 Redis）
# =========================
from pydantic import BaseModel

class TraceStep(BaseModel):
    step_type: str  # planner / executor / retriever / synthesizer
    timestamp: float
    data: dict

class TraceData(BaseModel):
    session_id: str
    question: str
    steps: List[TraceStep]
    start_time: float
    end_time: Optional[float] = None

class TraceStartRequest(BaseModel):
    session_id: str
    question: str

@app.post("/trace/start")
async def start_trace(req: TraceStartRequest):
    async with trace_lock:
        trace_store[req.session_id] = {
            "session_id": req.session_id,
            "question": req.question,
            "steps": [],
            "start_time": time.time(),
            "end_time": None
        }
    return {"status": "ok"}

@app.post("/trace/step")
async def add_trace_step(session_id: str, step_type: str, data: dict):
    """添加一个追踪步骤"""
    if session_id not in trace_store:
        return {"status": "error", "message": "trace not found"}
    
    async with trace_lock:
        trace_store[session_id]["steps"].append({
            "step_type": step_type,
            "timestamp": time.time(),
            "data": data
        })
    return {"status": "ok"}

@app.get("/trace/{session_id}")
async def get_trace(session_id: str):
    """获取完整的追踪链路"""
    if session_id not in trace_store:
        return {"success": False, "error": "trace not found"}
    
    trace = trace_store[session_id].copy()
    trace["steps"] = trace["steps"]
    return {"success": True, "data": trace}

@app.delete("/trace/{session_id}")
async def clear_trace(session_id: str):
    """清除追踪数据"""
    async with trace_lock:
        if session_id in trace_store:
            del trace_store[session_id]
    return {"status": "ok"}


## 审批助手侧边栏，支持待审批列表点击查看详情
class ApprovalDetailResponse(BaseModel):
    success: bool
    data: Optional[dict] = None
    error: Optional[str] = None

@app.get(
    "/approval/{approval_id}",
    tags=["审批"],
    summary="获取审批详情",
    description="根据审批 ID 获取完整的审批信息，包括 content 详情"
)
async def get_approval_detail(approval_id: str, request: Request):
    """
    获取指定审批项的完整详情
    """
    from tools.approval_tool import ApprovalTool
    from chat import get_current_tenant

    tenant_id = get_current_tenant()
    user = current_session_user if current_session_user else None

    if not user:
        return ApprovalDetailResponse(
            success=False,
            error="请先声明用户身份（用户：xxx）"
        )

    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute(
            """SELECT id, title, content, type, requester, requester_role,
                      reviewer, reviewer_role, status, comment, created_at, reviewed_at
               FROM approvals
               WHERE id = %s AND tenant_id = %s""",
            (approval_id, tenant_id)
        )
        row = cursor.fetchone()
        conn.close()

        if not row:
            return ApprovalDetailResponse(
                success=False,
                error=f"未找到审批项 {approval_id}"
            )

        # 权限校验：只有审批人或申请人才能查看详情
        if row.get('requester') != user and row.get('reviewer') != user:
            return ApprovalDetailResponse(
                success=False,
                error=f"您没有权限查看此审批项"
            )
        from utils.crypto import decrypt_if_needed
        content = decrypt_if_needed(row.get('content')) if row.get('content') else ""
        return ApprovalDetailResponse(
            success=True,
            data={
                "id": row.get('id'),
                "title": row.get('title'),
                "content": content,
                "type": row.get('type'),
                "requester": row.get('requester'),
                "requester_role": row.get('requester_role'),
                "reviewer": row.get('reviewer'),
                "reviewer_role": row.get('reviewer_role'),
                "status": row.get('status'),
                "comment": row.get('comment'),
                "created_at": row.get('created_at').strftime("%Y-%m-%d %H:%M:%S") if row.get('created_at') else "",
                "reviewed_at": row.get('reviewed_at').strftime("%Y-%m-%d %H:%M:%S") if row.get('reviewed_at') else ""
            }
        )
    except Exception as e:
        logger.error(f"获取审批详情失败: {e}")
        return ApprovalDetailResponse(
            success=False,
            error=str(e)
        )

@app.get(
    "/preview/{filename}",
    tags=["报告"],
    summary="预览评估表",
    description="在线预览评估表 Word 文档内容，无需下载"
)
async def preview_report(filename: str):
    """
    将 Word 文档内容转换为 HTML 预览
    """
    # 安全校验：只取文件名部分，防止路径遍历
    safe_filename = Path(filename).name
    file_path = (REPORTS_DIR / safe_filename).resolve()

    # 确保最终路径在 reports 目录内
    if not str(file_path).startswith(str(REPORTS_DIR)):
        return {"success": False, "error": "非法路径"}

    if not file_path.exists():
        return {"success": False, "error": "文件不存在"}
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        # ===== 提取段落内容 =====
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # ===== 提取表格中的所有文本 =====
        all_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in all_texts:
                        all_texts.append(text)
        
        # ===== 过滤无意义的单字 =====
        skip_words = ["用", "药", "计", "划", "单", "基", "本", "信", "息", "诊", "断", "评", "估", "结", "果", "方", "式"]
        filtered = []
        for t in all_texts:
            if len(t) >= 2 or t in ["用药计划单", "基本信息", "联系方式"]:
                filtered.append(t)
            elif t not in skip_words:
                filtered.append(t)
        
        # ===== 解析键值对 =====
        patient_info = {}
        fields_order = ["姓名", "性别", "年龄", "联系方式", "家庭住址", "身份证号", "临床诊断", "主要问题", "目前用药", "用药史", "过敏史", "症状描述", "评估结果", "用药目标", "用药注意事项"]
        
        # 提取键值对
        for text in filtered:
            for field in fields_order:
                if field in text:
                    # 提取值（去除字段名）
                    value = text.replace(field, "").replace("：", "").replace(":", "").strip()
                    if value:
                        patient_info[field] = value
                    break
            else:
                # 如果没有匹配到字段名，作为附加信息
                if "：" in text or ":" in text:
                    parts = re.split(r'[：:]', text, 1)
                    if len(parts) == 2:
                        key = parts[0].strip()
                        value = parts[1].strip()
                        patient_info[key] = value
                else:
                    # 尝试匹配常见格式：1. 诊断内容
                    match = re.match(r'^(\d+)\.\s*(.+)$', text)
                    if match:
                        patient_info[f"项目{match.group(1)}"] = match.group(2)
                    else:
                        patient_info["其他信息"] = text
        
        # ===== 构建美观的 HTML =====
        preview_html = """
        <html>
        <head>
            <meta charset="UTF-8">
            <title>评估表预览</title>
            <style>
                * { margin: 0; padding: 0; box-sizing: border-box; }
                body { 
                    font-family: 'Microsoft YaHei', 'PingFang SC', Arial, sans-serif; 
                    padding: 40px 20px; 
                    max-width: 900px; 
                    margin: 0 auto; 
                    background: #f0f4f8;
                    min-height: 100vh;
                }
                .container {
                    background: #ffffff;
                    border-radius: 16px;
                    box-shadow: 0 4px 24px rgba(0,0,0,0.08);
                    padding: 40px 48px;
                    margin-bottom: 20px;
                }
                .header {
                    text-align: center;
                    padding-bottom: 24px;
                    border-bottom: 3px solid #1a73e8;
                    margin-bottom: 28px;
                }
                .header h1 {
                    font-size: 28px;
                    color: #1a73e8;
                    font-weight: 700;
                    letter-spacing: 2px;
                }
                .header .subtitle {
                    color: #888;
                    font-size: 14px;
                    margin-top: 6px;
                }
                .info-grid {
                    display: grid;
                    grid-template-columns: 1fr 1fr;
                    gap: 12px 32px;
                    background: #f8f9fa;
                    border-radius: 10px;
                    padding: 16px 20px;
                    margin-bottom: 24px;
                }
                .info-grid .item {
                    display: flex;
                    align-items: baseline;
                    padding: 4px 0;
                }
                .info-grid .label {
                    font-weight: 600;
                    color: #555;
                    font-size: 14px;
                    min-width: 70px;
                }
                .info-grid .value {
                    color: #1a1a1a;
                    font-size: 14px;
                }
                
                .section {
                    margin-bottom: 20px;
                }
                .section-title {
                    font-size: 16px;
                    font-weight: 700;
                    color: #1a73e8;
                    padding-bottom: 6px;
                    border-bottom: 2px solid #e8ecf0;
                    margin-bottom: 12px;
                }
                .section-content {
                    padding: 8px 4px;
                    line-height: 1.8;
                    color: #333;
                    font-size: 14px;
                }
                .section-content .field {
                    padding: 4px 0;
                }
                .section-content .field-label {
                    font-weight: 600;
                    color: #555;
                }
                
                .risk-box {
                    background: #fef9e7;
                    border-left: 4px solid #f39c12;
                    padding: 16px 20px;
                    border-radius: 6px;
                    margin: 16px 0;
                    font-size: 14px;
                    line-height: 1.8;
                    color: #333;
                }
                .risk-box strong {
                    color: #e67e22;
                }
                
                .medication-table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 12px 0;
                    font-size: 14px;
                }
                .medication-table td {
                    padding: 10px 16px;
                    border-bottom: 1px solid #e8ecf0;
                    vertical-align: top;
                }
                .medication-table .label-cell {
                    font-weight: 600;
                    color: #555;
                    width: 100px;
                    background: #f8f9fa;
                }
                .medication-table .value-cell {
                    color: #333;
                }
                
                .footer {
                    text-align: center;
                    padding-top: 20px;
                    border-top: 1px solid #e8ecf0;
                    margin-top: 12px;
                }
                .footer .btn {
                    display: inline-block;
                    padding: 10px 28px;
                    background: #1a73e8;
                    color: #fff;
                    text-decoration: none;
                    border-radius: 8px;
                    font-size: 14px;
                    font-weight: 500;
                    transition: background 0.2s;
                    margin: 0 8px;
                }
                .footer .btn:hover {
                    background: #1557b0;
                }
                .footer .btn-outline {
                    background: #fff;
                    color: #555;
                    border: 1px solid #ddd;
                }
                .footer .btn-outline:hover {
                    background: #f5f5f5;
                }
                
                @media (max-width: 600px) {
                    .container { padding: 20px; }
                    .info-grid { grid-template-columns: 1fr; }
                }
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>📋 评估表预览</h1>
                    <div class="subtitle">用药方案评估报告</div>
                </div>
        """
        
        # ===== 基本信息卡片 =====
        basic_fields = ["姓名", "性别", "年龄", "联系方式", "家庭住址", "身份证号"]
        basic_items = []
        for field in basic_fields:
            if field in patient_info and patient_info[field]:
                basic_items.append(f'<div class="item"><span class="label">{field}</span><span class="value">{patient_info[field]}</span></div>')
        
        if basic_items:
            preview_html += '<div class="info-grid">' + "".join(basic_items) + '</div>'
        
        # ===== 段落内容 =====
        for p in paragraphs:
            if '风险告知书' in p:
                continue
            preview_html += f'<p style="font-size:14px;color:#333;line-height:1.8;margin:8px 0;">{p}</p>'
        
        # ===== 潜在风险告知书 =====
        risk_text = None
        for p in paragraphs:
            if '风险告知书' in p or '潜在意外风险' in p:
                risk_text = p
                break
        
        if risk_text:
            preview_html += f'''
            <div class="risk-box">
                <strong>⚠️ 潜在风险告知书</strong><br>
                {risk_text}
            </div>
            '''
        
        # ===== 用药计划单 =====
        preview_html += '''
                <div class="section">
                    <div class="section-title">💊 用药计划单</div>
                    <table class="medication-table">
        '''
        
        med_fields = ["临床诊断", "主要问题", "目前用药", "用药史", "过敏史", "症状描述", "评估结果", "用药目标", "用药注意事项"]
        for field in med_fields:
            if field in patient_info and patient_info[field]:
                preview_html += f'''
                        <tr>
                            <td class="label-cell">{field}</td>
                            <td class="value-cell">{patient_info[field]}</td>
                        </tr>
                '''
        
        preview_html += '''
                    </table>
                </div>
        '''
        
        # ===== 其他信息 =====
        other_items = []
        for key, value in patient_info.items():
            if key not in basic_fields and key not in med_fields:
                other_items.append(f'<div class="field"><span class="field-label">{key}：</span>{value}</div>')
        
        if other_items:
            preview_html += '''
                <div class="section">
                    <div class="section-title">📎 其他信息</div>
                    <div class="section-content">
            ''' + "".join(other_items) + '''
                    </div>
                </div>
            '''
        
        # ===== 底部按钮 =====
        preview_html += f'''
                <div class="footer">
                    <a href="/reports/{filename}" class="btn" download>📥 下载文档</a>
                    <a href="javascript:window.close()" class="btn btn-outline">✕ 关闭预览</a>
                </div>
            </div>
        </body>
        </html>
        '''
        
        return HTMLResponse(content=preview_html, media_type="text/html")
        
    except Exception as e:
        logger.error(f"预览评估表失败: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "error": f"预览失败: {str(e)}"}